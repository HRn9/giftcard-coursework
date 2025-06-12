from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os

def create_final_report():
    # Создаем новый документ
    doc = Document()
    
    # Настройка стилей
    styles = doc.styles
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    heading1_style = styles['Heading 1']
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    
    heading2_style = styles['Heading 2']
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True
    
    # Заголовок
    title = doc.add_heading('🎁 Gift Card Analytics System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Информация об авторе
    author_info = doc.add_paragraph()
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_info.add_run('Author: Mark Romanov\n').bold = True
    author_info.add_run('Group: 23-HO-6\n').bold = True
    author_info.add_run('Course Work Project\n').bold = True
    author_info.add_run('June 2025').bold = True
    
    doc.add_page_break()
    
    # Содержание
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. Technical Architecture',
        '3. Database Design',
        '4. ETL Process',
        '5. SQL Queries Analysis',
        '6. Power BI Dashboard',
        '7. Implementation Details',
        '8. Screenshots and Documentation',
        '9. Conclusion'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        'This course work project implements a complete data analytics solution for a gift card system. '
        'The project demonstrates the implementation of OLTP (Online Transaction Processing) and OLAP '
        '(Online Analytical Processing) databases, ETL (Extract, Transform, Load) processes, SQL queries, '
        'and Power BI visualizations.'
    )
    
    doc.add_heading('1.1 Project Objectives', level=2)
    objectives = [
        'Design and implement OLTP database with normalized structure',
        'Create OLAP data warehouse with star schema',
        'Implement ETL process for data migration',
        'Develop analytical SQL queries for both OLTP and OLAP',
        'Create interactive Power BI dashboard with visualizations',
        'Provide comprehensive documentation and screenshots'
    ]
    
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # 2. Technical Architecture
    doc.add_heading('2. Technical Architecture', level=1)
    doc.add_paragraph(
        'The project follows a modern data architecture pattern with separate OLTP and OLAP systems:'
    )
    
    doc.add_heading('2.1 System Components', level=2)
    components = [
        'OLTP Database (PostgreSQL): Normalized structure for transaction processing',
        'OLAP Data Warehouse (PostgreSQL): Denormalized structure for analytical queries',
        'ETL Process: Data extraction, transformation, and loading',
        'Power BI Dashboard: Interactive visualizations and reporting',
        'Documentation: Comprehensive project documentation and screenshots'
    ]
    
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')
    
    # 3. Database Design
    doc.add_heading('3. Database Design', level=1)
    
    doc.add_heading('3.1 OLTP Database (giftcards_oltp)', level=2)
    doc.add_paragraph(
        'The OLTP database follows Third Normal Form (3NF) normalization for efficient transaction processing. '
        'It contains 8 tables designed to handle day-to-day operations.'
    )
    
    oltp_tables = [
        'Users: Customer information and profiles',
        'GiftCardTypes: Different types of gift cards available',
        'GiftCards: Individual gift card instances',
        'TransactionTypes: Types of transactions (purchase, usage, etc.)',
        'MerchantCategories: Categories of merchants',
        'Merchants: Individual merchant information',
        'Transactions: Transaction records',
        'Promotions: Promotional campaigns and offers'
    ]
    
    for table in oltp_tables:
        doc.add_paragraph(table, style='List Bullet')
    
    doc.add_paragraph(
        'Reference: See screenshot "all_tables.png" for complete database structure visualization.'
    )
    
    doc.add_heading('3.2 OLAP Data Warehouse (giftcards_dwh)', level=2)
    doc.add_paragraph(
        'The OLAP data warehouse uses a star schema design optimized for analytical queries. '
        'It includes Slowly Changing Dimension (SCD) Type 2 implementation for historical tracking.'
    )
    
    doc.add_heading('Dimension Tables:', level=3)
    dim_tables = [
        'DimUser: User dimension with SCD Type 2',
        'DimGiftCard: Gift card dimension with SCD Type 2',
        'DimMerchant: Merchant dimension',
        'DimMerchantCategory: Merchant category dimension',
        'DimDate: Date dimension for time-based analysis',
        'DimPromotion: Promotion dimension'
    ]
    
    for dim in dim_tables:
        doc.add_paragraph(dim, style='List Bullet')
    
    doc.add_heading('Fact Tables:', level=3)
    fact_tables = [
        'FactCardSales: Sales transactions',
        'FactCardUsage: Usage transactions'
    ]
    
    for fact in fact_tables:
        doc.add_paragraph(fact, style='List Bullet')
    
    doc.add_heading('Bridge Table:', level=3)
    doc.add_paragraph('BridgeUserPromotion: Many-to-many relationship between users and promotions', style='List Bullet')
    
    doc.add_paragraph(
        'Reference: See screenshots "oltp.png" and "olap.png" for database schema diagrams.'
    )
    
    # 4. ETL Process
    doc.add_heading('4. ETL Process', level=1)
    doc.add_paragraph(
        'The ETL (Extract, Transform, Load) process is implemented to migrate data from the OLTP system '
        'to the OLAP data warehouse. The process includes data validation, transformation, and loading.'
    )
    
    doc.add_heading('4.1 ETL Components', level=2)
    etl_components = [
        'Data Loading (LoadData.psql): Initial data loading from CSV files to OLTP',
        'ETL Script (ETL.sql): Main ETL process for OLTP to OLAP migration',
        'Export Script (csv_to_excel.py): Data export for Power BI consumption'
    ]
    
    for comp in etl_components:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_heading('4.2 ETL Process Flow', level=2)
    etl_flow = [
        'Extract: Read data from OLTP tables',
        'Transform: Apply business rules and data transformations',
        'Load: Insert transformed data into OLAP tables',
        'Validate: Ensure data quality and integrity'
    ]
    
    for step in etl_flow:
        doc.add_paragraph(step, style='List Bullet')
    
    # 5. SQL Queries Analysis
    doc.add_heading('5. SQL Queries Analysis', level=1)
    
    doc.add_heading('5.1 OLTP Queries', level=2)
    doc.add_paragraph(
        'Three analytical queries were developed for the OLTP database to provide business insights '
        'from transactional data.'
    )
    
    oltp_queries = [
        'User Analysis: Customer spending patterns and transaction counts',
        'Sales Analysis: Gift card sales performance by time period',
        'Merchant Analysis: Transaction volume and revenue by merchant category'
    ]
    
    for query in oltp_queries:
        doc.add_paragraph(query, style='List Bullet')
    
    doc.add_paragraph(
        'Reference: See screenshot "sql_queries_oltp_results.png" for query execution results.'
    )
    
    doc.add_heading('5.2 OLAP Queries', level=2)
    doc.add_paragraph(
        'Three analytical queries were developed for the OLAP data warehouse to provide '
        'comprehensive business intelligence insights.'
    )
    
    olap_queries = [
        'Sales Performance: Monthly sales trends and growth rates',
        'Category Analysis: Usage patterns by merchant category',
        'Customer Segmentation: User behavior and spending analysis'
    ]
    
    for query in olap_queries:
        doc.add_paragraph(query, style='List Bullet')
    
    doc.add_paragraph(
        'Reference: See screenshot "sql_queries_olap_results.png" for query execution results.'
    )
    
    # 6. Power BI Dashboard
    doc.add_heading('6. Power BI Dashboard', level=1)
    doc.add_paragraph(
        'An interactive Power BI dashboard was created to provide visual insights into the gift card system. '
        'The dashboard includes multiple visualizations and interactive filtering capabilities.'
    )
    
    doc.add_heading('6.1 Dashboard Components', level=2)
    
    doc.add_heading('Title and Slicers:', level=3)
    slicers = [
        'Dashboard Title: "🎁 Gift Card Analytics Dashboard"',
        'Year Slicer: Filter data by year',
        'Merchant Category Slicer: Filter by merchant category'
    ]
    
    for slicer in slicers:
        doc.add_paragraph(slicer, style='List Bullet')
    
    doc.add_heading('KPI Cards:', level=3)
    kpi_cards = [
        'Total Sales: Revenue from gift card sales',
        'Total Usage: Amount spent using gift cards',
        'Active Cards: Number of active gift cards'
    ]
    
    for kpi in kpi_cards:
        doc.add_paragraph(kpi, style='List Bullet')
    
    doc.add_heading('Visualizations:', level=3)
    visualizations = [
        'Clustered Column Chart: Sales vs Usage by Month',
        'Pie Chart: Usage by Category',
        'Table: Top Spenders'
    ]
    
    for viz in visualizations:
        doc.add_paragraph(viz, style='List Bullet')
    
    doc.add_heading('6.2 Interactive Features', level=2)
    features = [
        'Cross-filtering between visualizations',
        'Dynamic filtering with slicers',
        'Responsive design for different screen sizes',
        'Professional color scheme and layout'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph(
        'Reference: See screenshot "powerbi_dashboard.png" for complete dashboard view.'
    )
    
    # 7. Implementation Details
    doc.add_heading('7. Implementation Details', level=1)
    
    doc.add_heading('7.1 Technology Stack', level=2)
    tech_stack = [
        'Database: PostgreSQL 12+',
        'ETL: SQL scripts and Python',
        'Visualization: Power BI Desktop/Web',
        'Programming: Python 3.8+ with pandas, psycopg2, openpyxl',
        'Documentation: Markdown and Word documents'
    ]
    
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('7.2 Project Structure', level=2)
    doc.add_paragraph(
        'The project is organized into logical directories for easy navigation and maintenance:'
    )
    
    structure = [
        'data/: Source CSV files (8 files with real data)',
        'sql/: SQL scripts for database creation, ETL, and queries',
        'scripts/: Python automation scripts',
        'documentation/: Screenshots, diagrams, and reports',
        'README.md: Comprehensive project documentation'
    ]
    
    for item in structure:
        doc.add_paragraph(item, style='List Bullet')
    
    # 8. Screenshots and Documentation
    doc.add_heading('8. Screenshots and Documentation', level=1)
    doc.add_paragraph(
        'Comprehensive documentation has been created to support the project implementation and demonstrate '
        'compliance with course work requirements.'
    )
    
    doc.add_heading('8.1 Screenshots', level=2)
    screenshots = [
        'all_tables.png: Complete database structure visualization',
        'sql_queries_oltp_results.png: OLTP query execution results',
        'sql_queries_olap_results.png: OLAP query execution results',
        'powerbi_dashboard.png: Power BI dashboard interface'
    ]
    
    for screenshot in screenshots:
        doc.add_paragraph(screenshot, style='List Bullet')
    
    doc.add_heading('8.2 Database Diagrams', level=2)
    diagrams = [
        'oltp.png: OLTP database schema diagram',
        'olap.png: OLAP data warehouse schema diagram'
    ]
    
    for diagram in diagrams:
        doc.add_paragraph(diagram, style='List Bullet')
    
    doc.add_heading('8.3 Technical Documentation', level=2)
    docs = [
        'README.md: Complete project overview and setup instructions',
        'SQL Scripts: Database creation, ETL, and query files',
        'Python Scripts: Data export and automation',
        'Power BI Files: Dashboard and data source files'
    ]
    
    for doc_item in docs:
        doc.add_paragraph(doc_item, style='List Bullet')
    
    # 9. Conclusion
    doc.add_heading('9. Conclusion', level=1)
    doc.add_paragraph(
        'This course work project successfully demonstrates the implementation of a complete data analytics '
        'solution for a gift card system. All requirements have been met and exceeded:'
    )
    
    achievements = [
        '✅ OLTP database with 8 tables in 3NF normalization',
        '✅ OLAP data warehouse with 9 tables in star schema',
        '✅ SCD Type 2 implementation for historical tracking',
        '✅ Bridge table for many-to-many relationships',
        '✅ Comprehensive ETL process for data migration',
        '✅ 3 OLTP queries for transactional analysis',
        '✅ 3 OLAP queries for analytical insights',
        '✅ Power BI dashboard with 6 visualizations and 2 slicers',
        '✅ Complete documentation with screenshots and diagrams',
        '✅ Professional project structure and implementation'
    ]
    
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_paragraph(
        'The project showcases modern data architecture principles and provides a solid foundation '
        'for business intelligence and analytics applications. The combination of OLTP and OLAP systems, '
        'coupled with interactive visualizations, creates a comprehensive solution for gift card analytics.'
    )
    
    # Сохраняем документ
    doc.save('documentation/reports/final_giftcards_db_report.docx')
    print("✅ Final report generated: documentation/reports/final_giftcards_db_report.docx")

if __name__ == "__main__":
    create_final_report() 